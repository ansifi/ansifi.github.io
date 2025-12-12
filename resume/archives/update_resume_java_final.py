#!/usr/bin/env python3
"""
Update resume with Java/Android project details - Final version with optimizations
"""

from docx import Document
from docx.shared import Pt

def update_resume_with_java_details():
    """Update resume with Java/Android project details"""
    
    input_file = "Ansif_Dev_Resume.docx"
    output_file = "Ansif_Dev_Resume.docx"
    
    doc = Document(input_file)
    
    # Track indices for insertion (work backwards to avoid index shifting)
    insertions = []
    modifications = []
    
    # First pass: identify what needs to be changed
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Update RedDotPayment project title to mention Antfarm/Reddotpayment
        if "1. RedDotPayment CRM System" in text:
            para.clear()
            run = para.add_run("1. RedDotPayment/Antfarm Android Payment Applications (Sept 2021 - April 2022)")
            run.font.size = Pt(10)
            run.bold = True
        
        # Update RedDotPayment Android description
        if "Developed comprehensive Android payment application using Java" in text and "QR code" in text:
            para.clear()
            run = para.add_run("Developed native Android payment applications using Java and Android SDK for RedDotPayment and Antfarm platforms")
            run.font.size = Pt(10)
        
        # Update Technologies for RedDotPayment - make more concise
        if "Technologies: Java, Android SDK, Codeigniter" in text:
            para.clear()
            run = para.add_run("Technologies: Java, Android SDK, Codeigniter, PHP, SQL, RESTful APIs, Payment Gateway Integration")
            run.font.size = Pt(10)
            # Mark position to add detailed bullets after this (condensed)
            insertions.append(('reddot_details', i + 1))
        
        # Mark position to add Java Spring Boot project
        if "Team size: 3" in text and i > 20:  # In CSR section
            # Check if next is empty or new client
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                if next_text == "" or next_text.startswith("Client:"):
                    insertions.append(('spring_boot', i + 1))
        
        # Mark position to add Android stock management
        if "Freelance Contracts: Jan 2013 – Sept 2016" in text:
            insertions.append(('android_stock', i + 1))
    
    # Second pass: insert new content (work backwards to preserve indices)
    insertions.sort(key=lambda x: x[1], reverse=True)
    
    for insert_type, index in insertions:
        if insert_type == 'reddot_details':
            # Add condensed Android payment app bullets
            bullets = [
                "• Built secure payment processing with encryption, tokenization, and QR code integration",
                "• Integrated RESTful APIs for payment gateway connectivity and transaction management"
            ]
            for bullet in reversed(bullets):
                new_para = doc.paragraphs[index].insert_paragraph_before()
                run = new_para.add_run(bullet)
                run.font.size = Pt(10)
        
        elif insert_type == 'spring_boot':
            # Add condensed Java Spring Boot microservices project
            content = [
                "",
                "4. Java Spring Boot Microservices Development (May 2022 - Sept 2024)",
                "Architected and developed multiple microservices using Java Spring Boot for enterprise applications",
                "Implemented RESTful APIs with Spring MVC, Spring Data JPA, and Hibernate for database operations",
                "Designed service communication using Spring Cloud, message queues, and implemented Spring Security with JWT",
                "Optimized performance with database query optimization, Redis caching, and API response improvements",
                "Technologies: Java, Spring Boot, Spring Cloud, Spring Data JPA, Spring Security, REST APIs, Microservices, MySQL, PostgreSQL, Redis",
                "Team size: 4-6"
            ]
            for item in reversed(content):
                new_para = doc.paragraphs[index].insert_paragraph_before()
                run = new_para.add_run(item)
                run.font.size = Pt(10)
                if item.startswith("4. Java"):
                    run.bold = True
        
        elif insert_type == 'android_stock':
            # Add condensed Android stock management apps
            content = [
                "Android Stock Management Applications (Jan 2012 - Dec 2014)",
                "Developed native Android apps for inventory management using Java and Android SDK",
                "Implemented SQLite database with synchronization, barcode scanning, and inventory reporting features",
                "Technologies: Java, Android SDK, SQLite, REST APIs, JSON parsing, Barcode Scanner"
            ]
            for item in reversed(content):
                new_para = doc.paragraphs[index].insert_paragraph_before()
                run = new_para.add_run(item)
                run.font.size = Pt(10)
                if item.startswith("Android Stock"):
                    run.bold = True
    
    # Save the document
    doc.save(output_file)
    print(f"Resume updated successfully: {output_file}")
    print("Added Java/Android project details:")
    print("  - Enhanced RedDotPayment/Antfarm Android payment app details")
    print("  - Added Java Spring Boot microservices project (2022-2024)")
    print("  - Added Android stock management apps (2012-2014)")
    print("  - Content optimized to fit within 7 pages")

if __name__ == "__main__":
    update_resume_with_java_details()


