#!/usr/bin/env python3
"""
Fix RedDotPayment section with proper Java/Android details
"""

from docx import Document
from docx.shared import Pt

def fix_reddot_section():
    """Fix RedDotPayment section"""
    
    input_file = "Ansif_Dev_Resume.docx"
    output_file = "Ansif_Dev_Resume.docx"
    
    doc = Document(input_file)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Update the Android description line
        if "Developed Android and web-based QR payment application" in text:
            para.clear()
            run = para.add_run("Developed native Android payment applications using Java and Android SDK for RedDotPayment and Antfarm platforms")
            run.font.size = Pt(10)
        
        # Update Technologies line
        if "Technologies: Codeigniter, PHP, SQL, Android Development" in text:
            para.clear()
            run = para.add_run("Technologies: Java, Android SDK, Codeigniter, PHP, SQL, RESTful APIs, Payment Gateway Integration")
            run.font.size = Pt(10)
            # Add detailed bullets after this
            bullets = [
                "• Built secure payment processing with encryption, tokenization, and QR code integration",
                "• Integrated RESTful APIs for payment gateway connectivity and transaction management"
            ]
            for bullet in reversed(bullets):
                new_para = doc.paragraphs[i + 1].insert_paragraph_before()
                run = new_para.add_run(bullet)
                run.font.size = Pt(10)
            break
    
    doc.save(output_file)
    print("RedDotPayment section updated successfully")

if __name__ == "__main__":
    fix_reddot_section()


