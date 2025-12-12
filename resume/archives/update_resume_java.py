#!/usr/bin/env python3
"""
Update resume with Java/Android project details
"""

from docx import Document
from docx.shared import Pt

def update_resume_with_java_details():
    """Update resume with Java/Android project details"""
    
    input_file = "Ansif_Dev_Resume.docx"
    output_file = "Ansif_Dev_Resume.docx"
    
    doc = Document(input_file)
    
    # Track indices for insertion (work backwards to avoid index shifting)
    insertions = []
    
    # First pass: identify what needs to be changed
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Update RedDotPayment Android description
        if "Developed Android and web-based QR payment application" in text:
            para.clear()
            run = para.add_run("Developed comprehensive Android payment application using Java with QR code payment capabilities")
            run.font.size = Pt(10)
        
        # Update Technologies for RedDotPayment
        if "Technologies: Codeigniter, PHP, SQL, Android Development" in text:
            para.clear()
            run = para.add_run("Technologies: Java, Android SDK, Codeigniter, PHP, SQL, RESTful APIs, Payment Gateway Integration, QR Code Processing")
            run.font.size = Pt(10)
            # Mark position to add detailed bullets after this
            insertions.append(('reddot_details', i + 1))
        
        # Mark position to add Java Spring Boot project
        if "Team size: 3" in text and i > 20:  # In CSR section
            # Check if next is empty or new client
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                if next_text == "" or next_text.startswith("Client:"):
                    insertions.append(('spring_boot', i + 1))
        
        # Mark position to add Android stock management
        if "Freelance Contracts: Jan 2013 – Sept 2016" in text:
            insertions.append(('android_stock', i + 1))
    
    # Second pass: insert new content (work backwards to preserve indices)
    insertions.sort(key=lambda x: x[1], reverse=True)
    
    for insert_type, index in insertions:
        if insert_type == 'reddot_details':
            # Add detailed Android payment app bullets
            bullets = [
                "• Developed native Android payment application using Java and Android SDK with secure payment processing",
                "• Implemented encryption, tokenization, and secure API integrations for payment gateway connectivity",
                "• Integrated QR code scanning and generation for seamless payment transactions"
            ]
            for bullet in reversed(bullets):
                new_para = doc.paragraphs[index].insert_paragraph_before()
                run = new_para.add_run(bullet)
                run.font.size = Pt(10)
        
        elif insert_type == 'spring_boot':
            # Add Java Spring Boot microservices project
            content = [
                "",
                "4. Java Spring Boot Microservices Development (May 2022 - Sept 2024)",
                "Architected and developed multiple microservices using Java Spring Boot framework for enterprise applications",
                "Implemented RESTful APIs with Spring MVC, Spring Data JPA, and Hibernate for efficient database operations",
                "Designed service-to-service communication patterns using Spring Cloud, message queues, and API gateways",
                "Implemented authentication and authorization using Spring Security, JWT tokens, and OAuth2",
                "Optimized database queries, implemented caching with Redis, and improved API response times",
                "Technologies: Java, Spring Boot, Spring Cloud, Spring Data JPA, Spring Security, REST APIs, Microservices Architecture, MySQL, PostgreSQL, Redis",
                "Team size: 4-6"
            ]
            for item in reversed(content):
                new_para = doc.paragraphs[index].insert_paragraph_before()
                run = new_para.add_run(item)
                run.font.size = Pt(10)
                if item.startswith("4. Java"):
                    run.bold = True
        
        elif insert_type == 'android_stock':
            # Add Android stock management apps
            content = [
                "Android Stock Management Applications (Jan 2012 - Dec 2014)",
                "Developed native Android applications for inventory and stock management using Java and Android SDK",
                "Implemented SQLite database for local data storage with synchronization capabilities",
                "Created intuitive user interfaces following Material Design principles and Android best practices",
                "Developed features for stock tracking, barcode scanning, inventory reporting, and data export",
                "Technologies: Java, Android SDK, SQLite, REST APIs, JSON parsing, Barcode Scanner Integration"
            ]
            for item in reversed(content):
                new_para = doc.paragraphs[index].insert_paragraph_before()
                run = new_para.add_run(item)
                run.font.size = Pt(10)
                if item.startswith("Android Stock"):
                    run.bold = True
    
    # Save the document
    doc.save(output_file)
    print(f"Resume updated successfully: {output_file}")
    print("Added Java/Android project details:")
    print("  - Enhanced RedDotPayment Android payment app details")
    print("  - Added Java Spring Boot microservices project (2022-2024)")
    print("  - Added Android stock management apps (2012-2014)")

if __name__ == "__main__":
    update_resume_with_java_details()
