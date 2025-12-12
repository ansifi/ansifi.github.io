#!/usr/bin/env python3
"""
Update resume with professional format including Sevendyne entrepreneurial experience
and remote/onsite employment projects
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_style(doc, text, level=1, bold=True, size=14):
    """Add formatted heading"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if level == 1:
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    para.paragraph_format.space_after = Pt(6)
    return para

def add_bullet_point(doc, text, indent=0):
    """Add formatted bullet point"""
    para = doc.add_paragraph()
    para.style = 'List Bullet'
    para.paragraph_format.left_indent = Inches(indent * 0.25)
    run = para.add_run(text)
    run.font.size = Pt(10)
    return para

def add_normal_text(doc, text, bold=False, italic=False):
    """Add normal formatted text"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = bold
    run.font.italic = italic
    para.paragraph_format.space_after = Pt(3)
    return para

def create_professional_resume():
    """Create professional resume document"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header Section
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name_run = header_para.add_run("ANSIF P I")
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Technical Consultant | Full Stack Developer | Entrepreneur | Startup Advisor")
    title_run.font.size = Pt(11)
    title_run.font.italic = True
    
    doc.add_paragraph()
    
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = "Email: anspi07@gmail.com | Phone: +91 95444 67629 | LinkedIn: linkedin.com/in/ansifpi | GitHub: github.com/ansifi"
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(9)
    
    portfolio_para = doc.add_paragraph()
    portfolio_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    portfolio_run = portfolio_para.add_run("Portfolio: 7dyne.com | Location: Kochi, Kerala, India")
    portfolio_run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Professional Summary
    add_heading_style(doc, "PROFESSIONAL SUMMARY", level=1, size=12)
    summary_text = ("Accomplished IT professional with 18+ years of comprehensive experience in software development, "
                   "technical consulting, and entrepreneurship. Proven expertise as a co-founder and technical leader, "
                   "specializing in full-stack development using Python, Java, React, and Node.js. Extensive experience "
                   "working with international clients in both remote and onsite capacities across multiple time zones. "
                   "Strong background in building scalable applications, leading technical teams, and managing complex "
                   "projects from conception to deployment. Founded Sevendyne Consultancy Services LLP in 2016, "
                   "successfully delivering 25+ projects and serving 10+ clients.")
    add_normal_text(doc, summary_text)
    
    doc.add_paragraph()
    
    # Core Technical Skills
    add_heading_style(doc, "CORE TECHNICAL SKILLS", level=1, size=12)
    
    skills = [
        ("Programming Languages:", "Python, Java, JavaScript, TypeScript, C++, C#, Ruby, PHP, Kotlin"),
        ("Backend Frameworks:", "FastAPI, Django, Flask, Spring Boot, Express.js, Ruby on Rails, Laravel, CodeIgniter"),
        ("Frontend Technologies:", "React, Angular, Node.js, HTML5, CSS3, Vue.js"),
        ("Mobile Development:", "Android (Java), React Native, Kotlin"),
        ("Databases:", "MySQL, PostgreSQL, SQL Server, MongoDB, SQLite, Redis"),
        ("DevOps & Cloud:", "AWS, Docker, Kubernetes, CI/CD Pipelines, Git, Jenkins"),
        ("AI & Machine Learning:", "OpenAI API, Machine Learning Algorithms, Data Analytics"),
        ("Architecture & Design:", "Microservices, RESTful APIs, System Architecture, Technical Due Diligence"),
        ("Tools & Methodologies:", "Agile Scrum, GitHub, JIRA, Asana, Trello, Bitbucket")
    ]
    
    for skill_type, skill_list in skills:
        skill_para = doc.add_paragraph()
        skill_type_run = skill_para.add_run(skill_type + " ")
        skill_type_run.font.size = Pt(10)
        skill_type_run.font.bold = True
        skill_list_run = skill_para.add_run(skill_list)
        skill_list_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Entrepreneurial Experience
    add_heading_style(doc, "ENTREPRENEURIAL EXPERIENCE", level=1, size=12)
    
    # Sevendyne Experience
    exp1_title = doc.add_paragraph()
    exp1_title_run = exp1_title.add_run("Co-Founder & Technical Consultant")
    exp1_title_run.font.size = Pt(11)
    exp1_title_run.font.bold = True
    
    exp1_company = doc.add_paragraph()
    exp1_company_run = exp1_company.add_run("Sevendyne Consultancy Services LLP | September 2016 – Present | Kochi, Kerala, India | Hybrid (Remote & Onsite)")
    exp1_company_run.font.size = Pt(10)
    exp1_company_run.font.italic = True
    
    exp1_desc = doc.add_paragraph()
    exp1_desc_run = exp1_desc.add_run("Entrepreneurial venture providing technical staffing, payroll services, and consulting solutions")
    exp1_desc_run.font.size = Pt(10)
    exp1_desc_run.font.italic = True
    
    doc.add_paragraph()
    
    exp1_key = doc.add_paragraph()
    exp1_key_run = exp1_key.add_run("Key Responsibilities:")
    exp1_key_run.font.size = Pt(10)
    exp1_key_run.font.bold = True
    
    responsibilities = [
        "Founded and co-manage Sevendyne Consultancy Services LLP, delivering technical solutions to global clients",
        "Lead technical teams (3-6 members) in developing enterprise-level applications",
        "Manage HR operations including payroll processing, employee onboarding, compliance, and documentation",
        "Provide technical consulting, architecture reviews, and startup advisory services",
        "Build and maintain client relationships across multiple time zones",
        "Oversee business development and client acquisition strategies"
    ]
    
    for resp in responsibilities:
        add_bullet_point(doc, resp, indent=1)
    
    doc.add_paragraph()
    
    exp1_ach = doc.add_paragraph()
    exp1_ach_run = exp1_ach.add_run("Key Achievements:")
    exp1_ach_run.font.size = Pt(10)
    exp1_ach_run.font.bold = True
    
    achievements = [
        "Successfully completed 10+ projects across diverse industries",
        "Served 5+ clients with recurring monthly contracts",
        "Built dedicated remote teams for international clients",
        "Established technical staffing and payroll services business model",
        "Currently managing multiple client projects simultaneously"
    ]
    
    for ach in achievements:
        add_bullet_point(doc, ach, indent=1)
    
    doc.add_paragraph()
    
    tech_para = doc.add_paragraph()
    tech_run = tech_para.add_run("Technologies Used: Python, Java, React, Node.js, Spring Boot, Django, PostgreSQL, MySQL, AWS, Docker")
    tech_run.font.size = Pt(9)
    tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # HRMS Project Development
    hrms_title = doc.add_paragraph()
    hrms_title_run = hrms_title.add_run("HRMS Project Development")
    hrms_title_run.font.size = Pt(11)
    hrms_title_run.font.bold = True
    
    hrms_duration = doc.add_paragraph()
    hrms_duration_run = hrms_duration.add_run("Sevendyne Consultancy Services LLP | January 2024 – November 2024 | Remote | Salary: ~₹1L/month")
    hrms_duration_run.font.size = Pt(10)
    hrms_duration_run.font.italic = True
    
    hrms_desc = doc.add_paragraph()
    hrms_desc_text = ("Developed comprehensive HRMS (Human Resource Management System) for Sevendyne startup using Django Python framework "
                     "and deployed on AWS cloud infrastructure. Led Python development and AWS deployment activities. "
                     "Built complete employee and client management portal with comprehensive HR features.")
    hrms_desc_run = hrms_desc.add_run(hrms_desc_text)
    hrms_desc_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    hrms_key = doc.add_paragraph()
    hrms_key_run = hrms_key.add_run("Key Features Implemented:")
    hrms_key_run.font.size = Pt(10)
    hrms_key_run.font.bold = True
    
    hrms_features = [
        "Employee Management Portal - Complete employee lifecycle management",
        "Client Management Portal - Client onboarding and relationship management",
        "Employee Onboarding - Automated onboarding workflows and documentation",
        "Payslip Generation - Automated payroll processing and payslip generation",
        "Attendance Management - Time tracking, attendance recording, and reporting",
        "Leave Management - Leave requests, approvals, balance tracking, and leave policies",
        "HR Policies & Compliance - Policy management and statutory compliance tracking",
        "Employee Self-Service Portal - Employee access to personal information and requests",
        "Reporting & Analytics - Comprehensive HR reports and analytics dashboard",
        "Role-based Access Control - Secure access management with different user roles"
    ]
    
    for feature in hrms_features:
        add_bullet_point(doc, feature, indent=1)
    
    doc.add_paragraph()
    
    hrms_tech = doc.add_paragraph()
    hrms_tech_run = hrms_tech.add_run("Technologies: Django, Python, AWS (EC2, RDS, S3, CloudFront), PostgreSQL, HTML/CSS, JavaScript, REST APIs")
    hrms_tech_run.font.size = Pt(9)
    hrms_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Project & Employment Experience
    add_heading_style(doc, "PROJECT & EMPLOYMENT EXPERIENCE", level=1, size=12)
    
    # CSR Informatik Projects
    exp2_title = doc.add_paragraph()
    exp2_title_run = exp2_title.add_run("Technical Consultant & Full Stack Developer")
    exp2_title_run.font.size = Pt(11)
    exp2_title_run.font.bold = True
    
    exp2_company = doc.add_paragraph()
    exp2_company_run = exp2_company.add_run("CSR Informatik GmbH, Germany | May 2022 – August 2024 | Remote")
    exp2_company_run.font.size = Pt(10)
    exp2_company_run.font.italic = True
    
    exp2_desc = doc.add_paragraph()
    exp2_desc_run = exp2_desc.add_run("Enterprise software development for German client - Multiple project engagements")
    exp2_desc_run.font.size = Pt(10)
    exp2_desc_run.font.italic = True
    
    doc.add_paragraph()
    
    # Project 1
    proj1_title = doc.add_paragraph()
    proj1_title_run = proj1_title.add_run("Project 1: Zoho Recruit Automation Platform")
    proj1_title_run.font.size = Pt(10)
    proj1_title_run.font.bold = True
    
    proj1_duration = doc.add_paragraph()
    proj1_duration_run = proj1_duration.add_run("Duration: June 2024 – August 2024 | Team Size: 4")
    proj1_duration_run.font.size = Pt(9)
    proj1_duration_run.font.italic = True
    
    proj1_points = [
        "Developed advanced Python-driven recruitment automation solution with OpenAI integration for intelligent job matching",
        "Implemented comprehensive search functionality with local SQL database management",
        "Created robust web application automating candidate screening and job posting workflows",
        "Integrated AI-powered candidate matching and automated recruitment processes"
    ]
    
    for point in proj1_points:
        add_bullet_point(doc, point, indent=1)
    
    proj1_tech = doc.add_paragraph()
    proj1_tech_run = proj1_tech.add_run("Technologies: Python, OpenAI API, SQL, Web Development, Automation, REST APIs")
    proj1_tech_run.font.size = Pt(9)
    proj1_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Project 2
    proj2_title = doc.add_paragraph()
    proj2_title_run = proj2_title.add_run("Project 2: E-commerce Platform with AI Integration")
    proj2_title_run.font.size = Pt(10)
    proj2_title_run.font.bold = True
    
    proj2_duration = doc.add_paragraph()
    proj2_duration_run = proj2_duration.add_run("Duration: June 2023 – June 2024 | Team Size: 4")
    proj2_duration_run.font.size = Pt(9)
    proj2_duration_run.font.italic = True
    
    proj2_points = [
        "Engineered custom e-commerce dashboard using Ruby on Rails with AI-powered web scraping capabilities",
        "Developed complex data extraction and transformation scripts using Python",
        "Implemented efficient product catalog management with SQL Server for data storage and analytics",
        "Built scalable backend architecture supporting high-volume transactions"
    ]
    
    for point in proj2_points:
        add_bullet_point(doc, point, indent=1)
    
    proj2_tech = doc.add_paragraph()
    proj2_tech_run = proj2_tech.add_run("Technologies: Ruby on Rails, Python, OpenAI, SQL Server, Web Scraping, REST APIs, Data Analytics")
    proj2_tech_run.font.size = Pt(9)
    proj2_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Project 3
    proj3_title = doc.add_paragraph()
    proj3_title_run = proj3_title.add_run("Project 3: Restaurant Analytics Dashboard")
    proj3_title_run.font.size = Pt(10)
    proj3_title_run.font.bold = True
    
    proj3_duration = doc.add_paragraph()
    proj3_duration_run = proj3_duration.add_run("Duration: May 2022 – June 2023 | Team Size: 3")
    proj3_duration_run.font.size = Pt(9)
    proj3_duration_run.font.italic = True
    
    proj3_points = [
        "Developed comprehensive data analytics dashboard with advanced visualizations using Apache eChart",
        "Created complex Angular frontend with Spring Boot backend",
        "Designed intricate SQL queries for real-time data aggregation, reporting, and business intelligence",
        "Implemented real-time data processing and visualization capabilities"
    ]
    
    for point in proj3_points:
        add_bullet_point(doc, point, indent=1)
    
    proj3_tech = doc.add_paragraph()
    proj3_tech_run = proj3_tech.add_run("Technologies: Angular, Spring Boot, SQL Server, Apache eChart, Data Visualization, REST APIs")
    proj3_tech_run.font.size = Pt(9)
    proj3_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Java Spring Boot Microservices Development
    exp3_title = doc.add_paragraph()
    exp3_title_run = exp3_title.add_run("Java Spring Boot Microservices Development")
    exp3_title_run.font.size = Pt(11)
    exp3_title_run.font.bold = True
    
    exp3_duration = doc.add_paragraph()
    exp3_duration_run = exp3_duration.add_run("Client Project | May 2022 – September 2024 | Remote")
    exp3_duration_run.font.size = Pt(10)
    exp3_duration_run.font.italic = True
    
    exp3_desc = doc.add_paragraph()
    exp3_desc_text = ("Architected and developed multiple microservices using Java Spring Boot for enterprise applications. "
                     "Implemented RESTful APIs with Spring MVC, Spring Data JPA, and Hibernate for database operations. "
                     "Designed service communication using Spring Cloud, message queues, and implemented Spring Security with JWT. "
                     "Optimized performance with database query optimization, Redis caching, and API response improvements. "
                     "Led technical team of 4-6 developers in microservices architecture implementation.")
    exp3_desc_run = exp3_desc.add_run(exp3_desc_text)
    exp3_desc_run.font.size = Pt(10)
    
    exp3_tech = doc.add_paragraph()
    exp3_tech_run = exp3_tech.add_run("Technologies: Java, Spring Boot, Spring Cloud, Spring Data JPA, Spring Security, REST APIs, Microservices, MySQL, PostgreSQL, Redis")
    exp3_tech_run.font.size = Pt(9)
    exp3_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Android Payment Applications
    exp4_title = doc.add_paragraph()
    exp4_title_run = exp4_title.add_run("Android Payment Applications Development")
    exp4_title_run.font.size = Pt(11)
    exp4_title_run.font.bold = True
    
    exp4_duration = doc.add_paragraph()
    exp4_duration_run = exp4_duration.add_run("RedDotPayment/Antfarm | September 2021 – April 2022 | Remote")
    exp4_duration_run.font.size = Pt(10)
    exp4_duration_run.font.italic = True
    
    exp4_desc = doc.add_paragraph()
    exp4_desc_text = ("Developed native Android payment applications using Java and Android SDK for RedDotPayment and Antfarm platforms. "
                     "Built secure payment processing with encryption, tokenization, and QR code integration. "
                     "Integrated RESTful APIs for payment gateway connectivity and transaction management. "
                     "Implemented comprehensive payment workflows and transaction processing systems.")
    exp4_desc_run = exp4_desc.add_run(exp4_desc_text)
    exp4_desc_run.font.size = Pt(10)
    
    exp4_tech = doc.add_paragraph()
    exp4_tech_run = exp4_tech.add_run("Technologies: Java, Android SDK, CodeIgniter, PHP, SQL, RESTful APIs, Payment Gateway Integration, QR Code Processing")
    exp4_tech_run.font.size = Pt(9)
    exp4_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Freelance Software Developer
    exp5_title = doc.add_paragraph()
    exp5_title_run = exp5_title.add_run("Freelance Software Developer")
    exp5_title_run.font.size = Pt(11)
    exp5_title_run.font.bold = True
    
    exp5_duration = doc.add_paragraph()
    exp5_duration_run = exp5_duration.add_run("Various Clients | January 2013 – September 2016 | Remote & Onsite")
    exp5_duration_run.font.size = Pt(10)
    exp5_duration_run.font.italic = True
    
    exp5_desc = doc.add_paragraph()
    exp5_desc_text = ("Delivered full-stack web applications using Python, Java, PHP, and JavaScript for diverse client requirements. "
                     "Developed mobile applications for Android platform. Provided technical consulting and architecture design services. "
                     "Worked with clients across different industries and time zones, delivering solutions on time and within budget.")
    exp5_desc_run = exp5_desc.add_run(exp5_desc_text)
    exp5_desc_run.font.size = Pt(10)
    
    exp5_tech = doc.add_paragraph()
    exp5_tech_run = exp5_tech.add_run("Technologies: Python, Java, PHP, JavaScript, Android, MySQL, PostgreSQL")
    exp5_tech_run.font.size = Pt(9)
    exp5_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Android Stock Management Applications
    exp6_title = doc.add_paragraph()
    exp6_title_run = exp6_title.add_run("Android Stock Management Applications")
    exp6_title_run.font.size = Pt(11)
    exp6_title_run.font.bold = True
    
    exp6_duration = doc.add_paragraph()
    exp6_duration_run = exp6_duration.add_run("Client Projects | January 2012 – December 2014 | Remote")
    exp6_duration_run.font.size = Pt(10)
    exp6_duration_run.font.italic = True
    
    exp6_desc = doc.add_paragraph()
    exp6_desc_text = ("Developed native Android apps for inventory management using Java and Android SDK. "
                     "Implemented SQLite database with synchronization, barcode scanning, and inventory reporting features. "
                     "Created offline-capable applications with cloud synchronization capabilities. "
                     "Built comprehensive inventory tracking and reporting systems for various business clients.")
    exp6_desc_run = exp6_desc.add_run(exp6_desc_text)
    exp6_desc_run.font.size = Pt(10)
    
    exp6_tech = doc.add_paragraph()
    exp6_tech_run = exp6_tech.add_run("Technologies: Java, Android SDK, SQLite, REST APIs, JSON parsing, Barcode Scanner")
    exp6_tech_run.font.size = Pt(9)
    exp6_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # QEHC, Saudi Arabia
    exp7_title = doc.add_paragraph()
    exp7_title_run = exp7_title.add_run("Software Developer")
    exp7_title_run.font.size = Pt(11)
    exp7_title_run.font.bold = True
    
    exp7_duration = doc.add_paragraph()
    exp7_duration_run = exp7_duration.add_run("QEHC, Saudi Arabia | January 2012 – December 2012 | Onsite")
    exp7_duration_run.font.size = Pt(10)
    exp7_duration_run.font.italic = True
    
    exp7_desc = doc.add_paragraph()
    exp7_desc_text = ("Developed healthcare management applications using modern web technologies for healthcare data management. "
                     "Worked on cross-platform solutions for healthcare data management and patient information systems. "
                     "Collaborated with international team on enterprise-level healthcare projects. "
                     "Implemented secure data handling and compliance with healthcare regulations.")
    exp7_desc_run = exp7_desc.add_run(exp7_desc_text)
    exp7_desc_run.font.size = Pt(10)
    
    exp7_tech = doc.add_paragraph()
    exp7_tech_run = exp7_tech.add_run("Technologies: Web Technologies, Database Systems, Healthcare Systems")
    exp7_tech_run.font.size = Pt(9)
    exp7_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Hexosys SDN BHD, Malaysia
    exp8_title = doc.add_paragraph()
    exp8_title_run = exp8_title.add_run("Junior Software Developer")
    exp8_title_run.font.size = Pt(11)
    exp8_title_run.font.bold = True
    
    exp8_duration = doc.add_paragraph()
    exp8_duration_run = exp8_duration.add_run("Hexosys SDN BHD, Malaysia | October 2009 – October 2011 | Onsite")
    exp8_duration_run.font.size = Pt(10)
    exp8_duration_run.font.italic = True
    
    exp8_desc = doc.add_paragraph()
    exp8_desc_text = ("Developed C++ Qt and VC++ based USB protocol testing applications for embedded systems and industrial applications. "
                     "Debugged USB device protocols including TCP/UDP communication and device driver integration. "
                     "Worked on embedded systems and industrial automation projects. "
                     "Collaborated with cross-functional teams of 10+ members on complex hardware-software integration projects.")
    exp8_desc_run = exp8_desc.add_run(exp8_desc_text)
    exp8_desc_run.font.size = Pt(10)
    
    exp8_tech = doc.add_paragraph()
    exp8_tech_run = exp8_tech.add_run("Technologies: C++, Qt, VC++, USB Protocols, TCP/UDP, Embedded Systems")
    exp8_tech_run.font.size = Pt(9)
    exp8_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Integral & UCI Tech, India
    exp9_title = doc.add_paragraph()
    exp9_title_run = exp9_title.add_run("Trainee Software Engineer")
    exp9_title_run.font.size = Pt(11)
    exp9_title_run.font.bold = True
    
    exp9_duration = doc.add_paragraph()
    exp9_duration_run = exp9_duration.add_run("Integral & UCI Tech, India | April 2007 – April 2009 | Onsite")
    exp9_duration_run.font.size = Pt(10)
    exp9_duration_run.font.italic = True
    
    exp9_desc = doc.add_paragraph()
    exp9_desc_text = ("Developed trading room software using .NET C# and C++ frameworks for financial trading systems. "
                     "Integrated .NET SDK for external connections and real-time data processing. "
                     "Assisted senior developers in building applications and troubleshooting code issues. "
                     "Gained foundational experience in software development lifecycle, best practices, and financial software development.")
    exp9_desc_run = exp9_desc.add_run(exp9_desc_text)
    exp9_desc_run.font.size = Pt(10)
    
    exp9_tech = doc.add_paragraph()
    exp9_tech_run = exp9_tech.add_run("Technologies: C#, .NET, C++, Trading Systems, Financial Software")
    exp9_tech_run.font.size = Pt(9)
    exp9_tech_run.font.italic = True
    
    doc.add_paragraph()
    
    # Key Achievements Section
    add_heading_style(doc, "KEY ACHIEVEMENTS", level=1, size=12)
    
    key_achievements = [
        "18+ years of hands-on software development experience",
        "Founded Sevendyne Consultancy Services LLP (September 2016)",
        "25+ projects completed across various domains and industries",
        "10+ contracts served with recurring client relationships",
        "Active GitHub contributor with production applications",
        "Successfully led teams of 3-6 developers on multiple projects",
        "Extensive experience with remote and onsite project delivery",
        "Proven track record in entrepreneurship and business development"
    ]
    
    for achievement in key_achievements:
        add_bullet_point(doc, achievement)
    
    doc.add_paragraph()
    
    # Save document
    output_file = "Ansif_Dev_Resume_Professional.docx"
    doc.save(output_file)
    print(f"Professional resume created: {output_file}")

if __name__ == "__main__":
    create_professional_resume()

